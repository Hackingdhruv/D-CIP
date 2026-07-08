"""Report export service — PDF, DOCX, HTML, and JSON generation.

PDF: extends the existing dependency-free manual writer with richer formatting
     (headings, sections, tables, cover page, page numbers, headers/footers).
DOCX: uses the already-installed python-docx library.
HTML: pure Python string generation — no template engine required.
JSON: direct serialisation of sections_content.

Every export clearly labels AI-generated content.
"""

from __future__ import annotations

import hashlib
import json
from datetime import datetime, timezone
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from app.models.report import InvestigationReport

# ── Helpers ───────────────────────────────────────────────────────────────────

def _fmt_ts(ts_str: str | None) -> str:
    if not ts_str:
        return "N/A"
    try:
        dt = datetime.fromisoformat(ts_str.replace("Z", "+00:00"))
        return dt.strftime("%Y-%m-%d %H:%M UTC")
    except Exception:
        return str(ts_str)


def _file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


# ── JSON export ───────────────────────────────────────────────────────────────

def export_json(report: "InvestigationReport") -> bytes:
    payload = {
        "report_id": str(report.id),
        "case_id": str(report.case_id),
        "title": report.title,
        "report_type": report.report_type,
        "template": report.template,
        "version": report.version,
        "status": report.status,
        "content_hash": report.content_hash,
        "generated_at": report.generated_at.isoformat() if report.generated_at else None,
        "sections_content": report.sections_content,
    }
    return json.dumps(payload, indent=2, default=str).encode("utf-8")


# ── HTML export ───────────────────────────────────────────────────────────────

def export_html(report: "InvestigationReport") -> bytes:
    content = report.sections_content
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    cover = content.get("cover", {})
    title = cover.get("case_title") or report.title
    classification = cover.get("classification", "")
    generated_by = cover.get("generated_by", "")

    body_parts: list[str] = []

    for sec_type, sec in content.items():
        if not isinstance(sec, dict):
            continue
        sec_title = sec.get("title", sec_type)
        ai_flag = ""
        if sec.get("is_ai_generated"):
            ai_flag = (
                '<p class="ai-disclaimer">⚠️ AI-Generated Content — '
                "must be independently verified by the lead investigator.</p>"
            )

        if sec_type == "cover":
            continue  # rendered in header

        body_parts.append(f'<section id="{sec_type}">')
        body_parts.append(f"<h2>{_h(sec_title)}</h2>")
        if ai_flag:
            body_parts.append(ai_flag)

        if sec_type == "executive_summary":
            if sec.get("summary_text"):
                body_parts.append(f'<p class="summary">{_h(sec["summary_text"])}</p>')
            _html_list(body_parts, "Key Findings", sec.get("key_findings", []))
            _html_list(body_parts, "Potential Leads", sec.get("potential_leads", []))
            _html_list(body_parts, "Open Questions", sec.get("open_questions", []))

        elif sec_type == "case_overview":
            rows = [
                ("Reference", sec.get("reference_number")),
                ("Status", sec.get("status")),
                ("Priority", sec.get("priority")),
                ("Category", sec.get("category")),
                ("Created", _fmt_ts(sec.get("created_at"))),
            ]
            body_parts.append("<table>")
            for k, v in rows:
                if v:
                    body_parts.append(f"<tr><th>{k}</th><td>{_h(str(v))}</td></tr>")
            body_parts.append("</table>")
            if sec.get("description"):
                body_parts.append(f'<p>{_h(sec["description"])}</p>')
            team = sec.get("team", [])
            if team:
                body_parts.append("<h3>Investigation Team</h3><ul>")
                for m in team:
                    body_parts.append(f'<li>{_h(m["name"])} — {_h(m["role"])}</li>')
                body_parts.append("</ul>")

        elif sec_type == "evidence_inventory":
            items = sec.get("items", [])
            body_parts.append(f"<p>Total: {sec.get('total_count', len(items))} items</p>")
            if items:
                body_parts.append(
                    "<table><tr><th>File</th><th>Type</th><th>Size (bytes)</th>"
                    "<th>Uploaded By</th><th>Date</th><th>Status</th><th>Entities</th></tr>"
                )
                for item in items:
                    body_parts.append(
                        f"<tr>"
                        f"<td>{_h(item.get('filename',''))}</td>"
                        f"<td>{_h(item.get('mime_type',''))}</td>"
                        f"<td>{item.get('file_size','')}</td>"
                        f"<td>{_h(item.get('uploaded_by',''))}</td>"
                        f"<td>{_fmt_ts(item.get('uploaded_at'))}</td>"
                        f"<td>{_h(item.get('status',''))}</td>"
                        f"<td>{item.get('entity_count',0)}</td>"
                        f"</tr>"
                    )
                body_parts.append("</table>")

        elif sec_type == "timeline":
            items = sec.get("items", [])
            body_parts.append(f"<p>Total: {len(items)} events</p>")
            if items:
                body_parts.append("<div class='timeline'>")
                for item in items:
                    ts = _fmt_ts(item.get("event_timestamp"))
                    body_parts.append(
                        f'<div class="event">'
                        f'<span class="ts">{ts}</span>'
                        f'<strong>{_h(item.get("title",""))}</strong>'
                        f'<span class="badge">{_h(item.get("event_type",""))}</span>'
                    )
                    if item.get("description"):
                        body_parts.append(f'<p>{_h(item["description"])}</p>')
                    body_parts.append("</div>")
                body_parts.append("</div>")

        elif sec_type == "entities":
            by_type = sec.get("by_type", {})
            for etype, ents in by_type.items():
                if not ents:
                    continue
                body_parts.append(f"<h3>{_h(etype.replace('_',' ').title())}</h3><ul>")
                for ent in ents:
                    ctx = f" — {_h(ent['context'])}" if ent.get("context") else ""
                    conf = f" ({int(ent.get('confidence',1)*100)}%)" if ent.get("confidence") else ""
                    body_parts.append(f'<li>{_h(ent["value"])}{conf}{ctx}</li>')
                body_parts.append("</ul>")

        elif sec_type == "ai_findings":
            disc = sec.get("disclaimer")
            if disc:
                body_parts.append(f'<p class="ai-disclaimer">{_h(disc)}</p>')
            cs = sec.get("case_summary", {})
            if cs.get("summary_text"):
                body_parts.append(f"<h3>Case Analysis</h3><p>{_h(cs['summary_text'])}</p>")
            _html_list(body_parts, "Key Findings", cs.get("key_findings", []))
            _html_list(body_parts, "Potential Leads", cs.get("potential_leads", []))
            for es in sec.get("evidence_summaries", []):
                fname = _h(es.get("filename", ""))
                body_parts.append(f"<h3>Evidence: {fname}</h3>")
                if es.get("summary_text"):
                    body_parts.append(f'<p>{_h(es["summary_text"])}</p>')

        elif sec_type == "notes_tasks":
            notes = sec.get("notes", [])
            tasks = sec.get("tasks", [])
            if notes:
                body_parts.append("<h3>Notes</h3>")
                for n in notes:
                    pin = " 📌" if n.get("is_pinned") else ""
                    body_parts.append(
                        f'<div class="note"><strong>{_h(n["title"])}{pin}</strong>'
                        f'<small>{_h(n.get("created_by",""))} · {_fmt_ts(n.get("created_at"))}</small>'
                        f'<p>{_h(n.get("content",""))}</p></div>'
                    )
            if tasks:
                body_parts.append(
                    "<h3>Tasks</h3><table>"
                    "<tr><th>Title</th><th>Status</th><th>Priority</th><th>Due</th></tr>"
                )
                for t in tasks:
                    body_parts.append(
                        f"<tr><td>{_h(t['title'])}</td>"
                        f"<td>{_h(t['status'])}</td>"
                        f"<td>{_h(t['priority'])}</td>"
                        f"<td>{_fmt_ts(t.get('due_date'))}</td></tr>"
                    )
                body_parts.append("</table>")

        elif sec_type == "chain_of_custody":
            for item in sec.get("items", []):
                body_parts.append(f'<h3>{_h(item["filename"])}</h3>')
                body_parts.append(f'<p class="mono">SHA-256: {_h(item.get("sha256_hash","N/A"))}</p>')
                evts = item.get("events", [])
                if evts:
                    body_parts.append("<table><tr><th>Action</th><th>Actor</th><th>Description</th><th>Time</th></tr>")
                    for ev in evts:
                        body_parts.append(
                            f"<tr><td>{_h(ev['action'])}</td>"
                            f"<td>{_h(ev['actor'])}</td>"
                            f"<td>{_h(ev.get('description',''))}</td>"
                            f"<td>{_fmt_ts(ev.get('timestamp'))}</td></tr>"
                        )
                    body_parts.append("</table>")

        body_parts.append("</section>")

    body = "\n".join(body_parts)

    html = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{_h(title)} — Investigation Report</title>
<style>
  body {{font-family:'Segoe UI',Arial,sans-serif;margin:0;color:#1a1a2e;background:#fff}}
  .cover {{background:#1a1a2e;color:#fff;padding:60px 80px;min-height:200px}}
  .cover h1 {{font-size:2.2em;margin:0 0 8px}}
  .cover .sub {{opacity:.7;font-size:1.1em}}
  .classification {{
    display:inline-block;border:2px solid currentColor;
    padding:4px 16px;font-weight:700;letter-spacing:.15em;
    margin-bottom:24px;font-size:.9em
  }}
  .meta-bar {{background:#f0f0f8;padding:12px 80px;font-size:.85em;color:#555;
    border-bottom:1px solid #dde}}
  main {{max-width:960px;margin:0 auto;padding:40px 80px}}
  section {{margin-bottom:48px;border-bottom:1px solid #eee;padding-bottom:32px}}
  h2 {{font-size:1.5em;color:#1a1a2e;border-left:4px solid #4f46e5;
    padding-left:12px;margin-top:0}}
  h3 {{font-size:1.1em;color:#374151;margin-top:24px}}
  table {{border-collapse:collapse;width:100%;margin:12px 0;font-size:.9em}}
  th {{background:#f3f4f6;text-align:left;padding:8px 12px;border:1px solid #ddd}}
  td {{padding:7px 12px;border:1px solid #ddd}}
  tr:nth-child(even) td {{background:#fafafa}}
  .ai-disclaimer {{background:#fef3c7;border-left:4px solid #f59e0b;
    padding:10px 16px;margin:12px 0;font-size:.85em;color:#92400e}}
  .timeline .event {{
    border-left:3px solid #4f46e5;padding:8px 16px;margin:8px 0;
    background:#f9fafb;border-radius:0 6px 6px 0
  }}
  .event .ts {{color:#6b7280;font-size:.85em;display:block;margin-bottom:4px}}
  .event .badge {{
    background:#e0e7ff;color:#4338ca;font-size:.75em;
    padding:2px 8px;border-radius:99px;margin-left:8px
  }}
  .note {{border:1px solid #e5e7eb;padding:12px 16px;border-radius:6px;margin:8px 0}}
  .note small {{color:#6b7280;display:block;margin:4px 0 8px}}
  .summary {{font-size:1.05em;line-height:1.7;color:#374151}}
  .mono {{font-family:monospace;font-size:.8em;color:#6b7280;word-break:break-all}}
  ul {{padding-left:24px;line-height:1.8}}
  @media print {{
    .cover {{-webkit-print-color-adjust:exact;print-color-adjust:exact}}
    section {{page-break-inside:avoid}}
  }}
</style>
</head>
<body>
<div class="cover">
  <div class="classification">{_h(classification)}</div>
  <h1>{_h(title)}</h1>
  <p class="sub">
    {_h(cover.get("case_reference",""))} · v{report.version} · Generated {now}
    {(" by " + _h(generated_by)) if generated_by else ""}
  </p>
</div>
<div class="meta-bar">
  Report type: <strong>{_h(report.report_type)}</strong> &nbsp;|&nbsp;
  Template: <strong>{_h(report.template)}</strong> &nbsp;|&nbsp;
  SHA-256: <code>{report.content_hash or "N/A"}</code>
</div>
<main>
{body}
</main>
</body>
</html>"""

    return html.encode("utf-8")


def _h(s: str | None) -> str:
    if not s:
        return ""
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _html_list(parts: list[str], heading: str, items: list) -> None:
    if not items:
        return
    parts.append(f"<h3>{_h(heading)}</h3><ul>")
    for it in items:
        parts.append(f"<li>{_h(str(it))}</li>")
    parts.append("</ul>")


# ── PDF export (enhanced manual writer) ──────────────────────────────────────

_PW = 612   # US Letter points
_PH = 792
_ML = 60    # left margin
_MR = 60    # right margin
_MT = 60    # top margin
_MB = 60    # bottom margin
_FONT_NORMAL = 9
_FONT_HEADING = 13
_FONT_TITLE = 18
_LEADING_NORMAL = 13
_LEADING_HEADING = 20
_CHARS_NORMAL = 100
_CHARS_HEADING = 80

# Keep the simple char-based approach since we have no font metrics library.
# We simulate heading style via larger font-size commands in the content stream.


def export_pdf(report: "InvestigationReport") -> bytes:
    content = report.sections_content
    cover = content.get("cover", {})
    case_ref = cover.get("case_reference", "")
    classification = cover.get("classification", "")
    generated_by = cover.get("generated_by", "")
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Build logical lines
    all_lines: list[tuple[int, str]] = []  # (font_size, text)

    def h(text: str | None, size: int = _FONT_NORMAL) -> None:
        if not text:
            return
        max_chars = max(20, int(_CHARS_NORMAL * _FONT_NORMAL / size))
        for chunk in _chunk(str(text), max_chars):
            all_lines.append((size, chunk))

    def blank() -> None:
        all_lines.append((_FONT_NORMAL, ""))

    # Cover section
    h(classification, 11)
    h(cover.get("case_title") or report.title, _FONT_TITLE)
    blank()
    h(f"Case Reference: {case_ref}", 10)
    h(f"Report Type:    {report.report_type.replace('_',' ').title()}", 10)
    h(f"Version:        {report.version}", 10)
    h(f"Generated:      {now_str}", 10)
    if generated_by:
        h(f"Generated by:   {generated_by}", 10)
    h("-" * _CHARS_NORMAL)
    blank()
    blank()

    for sec_key, sec in content.items():
        if not isinstance(sec, dict) or sec_key == "cover":
            continue
        sec_title = sec.get("title", sec_key)
        h(sec_title.upper(), _FONT_HEADING)
        if sec.get("is_ai_generated"):
            h("[AI-GENERATED: verify independently]", 8)
        blank()
        _pdf_section(all_lines, sec)
        blank()
        h("-" * _CHARS_NORMAL)
        blank()

    # Footer note
    blank()
    h(f"Report SHA-256: {report.content_hash or 'N/A'}", 8)
    h(f"D-CIP Investigation Platform · {now_str}", 8)

    # Render PDF with page numbers
    return _render_pdf_pages(all_lines, case_ref, classification)


def _pdf_section(lines: list, sec: dict) -> None:
    def h(text: str | None, size: int = _FONT_NORMAL) -> None:
        if not text:
            return
        max_c = max(20, int(_CHARS_NORMAL * _FONT_NORMAL / size))
        for chunk in _chunk(str(text), max_c):
            lines.append((size, chunk))

    def blank() -> None:
        lines.append((_FONT_NORMAL, ""))

    sec_type = sec.get("type", "")

    if sec_type == "table_of_contents":
        h("(Table of contents generated at time of export)")

    elif sec_type == "executive_summary":
        if sec.get("summary_text"):
            h(sec["summary_text"])
            blank()
        for heading, key in [("Key Findings", "key_findings"), ("Potential Leads", "potential_leads"), ("Open Questions", "open_questions")]:
            items = sec.get(key, [])
            if items:
                h(f"  {heading}:", 10)
                for item in items:
                    h(f"    • {item}")
                blank()

    elif sec_type == "case_overview":
        for k, v in [
            ("Reference", sec.get("reference_number")),
            ("Status", sec.get("status")),
            ("Priority", sec.get("priority")),
            ("Category", sec.get("category")),
            ("Evidence Files", sec.get("evidence_count")),
            ("Tasks", sec.get("task_count")),
            ("Notes", sec.get("note_count")),
            ("Created", _fmt_ts(sec.get("created_at"))),
        ]:
            if v is not None:
                h(f"  {k}: {v}")
        if sec.get("description"):
            blank()
            h(sec["description"])
        team = sec.get("team", [])
        if team:
            blank()
            h("  Team:")
            for m in team:
                h(f"    • {m['name']} ({m['role']})")

    elif sec_type == "evidence_inventory":
        h(f"  Total items: {sec.get('total_count', 0)}")
        blank()
        for item in sec.get("items", []):
            h(f"  ▶ {item.get('filename', '')}", 10)
            h(f"    Type: {item.get('mime_type','')}  Size: {item.get('file_size','')} bytes")
            h(f"    Status: {item.get('status','')}  Uploaded by: {item.get('uploaded_by','')}")
            h(f"    SHA-256: {item.get('sha256_hash','N/A')}", 8)
            blank()

    elif sec_type == "timeline":
        h(f"  Total events: {sec.get('total_count', 0)}")
        if sec.get("date_from"):
            h(f"  From: {_fmt_ts(sec['date_from'])}  To: {_fmt_ts(sec.get('date_to'))}")
        blank()
        for item in sec.get("items", []):
            ts = _fmt_ts(item.get("event_timestamp"))
            h(f"  [{ts}] {item.get('title','')}", 10)
            h(f"    Type: {item.get('event_type','')}  Confidence: {int(item.get('confidence',0)*100)}%")
            if item.get("description"):
                h(f"    {item['description']}")
            blank()

    elif sec_type == "entities":
        h(f"  Total entities: {sec.get('total_count', 0)}")
        blank()
        for etype, ents in sec.get("by_type", {}).items():
            if not ents:
                continue
            h(f"  {etype.replace('_',' ').upper()} ({len(ents)}):", 10)
            for ent in ents[:20]:
                conf = f"{int(ent.get('confidence',1)*100)}%"
                h(f"    • {ent['value']}  [{conf}]")
            blank()

    elif sec_type == "ai_findings":
        disc = sec.get("disclaimer")
        if disc:
            h(f"  *** {disc} ***", 8)
            blank()
        cs = sec.get("case_summary", {})
        if cs.get("summary_text"):
            h("  Case Analysis:")
            h(cs["summary_text"])
            blank()
        for heading, key in [("Key Findings", "key_findings"), ("Potential Leads", "potential_leads")]:
            items = cs.get(key, [])
            if items:
                h(f"  {heading}:")
                for it in items:
                    h(f"    • {it}")
                blank()
        for es in sec.get("evidence_summaries", []):
            h(f"  Evidence: {es.get('filename', '')}", 10)
            if es.get("summary_text"):
                h(f"  {es['summary_text']}")
            blank()

    elif sec_type == "notes_tasks":
        notes = sec.get("notes", [])
        tasks = sec.get("tasks", [])
        if notes:
            h(f"  Notes ({len(notes)}):", 10)
            for n in notes:
                pin = " [PINNED]" if n.get("is_pinned") else ""
                h(f"    ▶ {n['title']}{pin}")
                h(f"      By: {n.get('created_by','')} on {_fmt_ts(n.get('created_at'))}", 8)
                if n.get("content"):
                    h(f"      {n['content']}")
                blank()
        if tasks:
            h(f"  Tasks ({len(tasks)}):", 10)
            for t in tasks:
                h(f"    [{t.get('status','').upper()}] {t['title']}")
                h(f"    Priority: {t.get('priority','')}  Due: {_fmt_ts(t.get('due_date'))}", 8)
                blank()

    elif sec_type == "chain_of_custody":
        h(f"  Evidence items: {sec.get('total_evidence', 0)}")
        blank()
        for item in sec.get("items", []):
            h(f"  ▶ {item.get('filename', '')}", 10)
            h(f"    SHA-256: {item.get('sha256_hash','N/A')}", 8)
            for ev in item.get("events", []):
                h(f"    [{_fmt_ts(ev.get('timestamp'))}] {ev.get('action','').upper()}"
                  f" — {ev.get('actor','')}  {ev.get('description','')}")
            blank()


def _chunk(text: str, max_chars: int) -> list[str]:
    text = text.replace("\t", "    ").replace("\r\n", " ").replace("\n", " ").strip()
    if not text:
        return [""]
    out = []
    while len(text) > max_chars:
        # Try to break on a space
        idx = text.rfind(" ", 0, max_chars)
        if idx < 0:
            idx = max_chars
        out.append(text[:idx])
        text = text[idx:].lstrip()
    out.append(text)
    return out


def _esc_pdf(text: str) -> str:
    return (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\n", " ")
    )


def _render_pdf_pages(
    all_lines: list[tuple[int, str]],
    case_ref: str,
    classification: str,
) -> bytes:
    """Render the logical lines into a multi-page PDF with headers/footers."""
    # Compute available vertical space per page (excluding header + footer)
    HEADER_H = 24
    FOOTER_H = 24
    TOP_Y = _PH - _MT - HEADER_H - 4
    BOTTOM_Y = _MB + FOOTER_H + 4

    objects: list[bytes] = []
    page_content_ids: list[int] = []
    page_ids: list[int] = []

    def add_obj(body: bytes) -> int:
        objects.append(body)
        return len(objects)

    font_id = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Courier >>")
    bold_id = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Courier-Bold >>")

    # Paginate: calculate how many lines fit per page
    def line_height(size: int) -> float:
        return size * 1.4

    pages: list[list[tuple[int, str]]] = []
    current_page: list[tuple[int, str]] = []
    current_y = TOP_Y

    for size, text in all_lines:
        lh = line_height(size)
        if current_y - lh < BOTTOM_Y:
            pages.append(current_page)
            current_page = []
            current_y = TOP_Y
        current_page.append((size, text))
        current_y -= lh

    if current_page:
        pages.append(current_page)
    if not pages:
        pages = [[(_FONT_NORMAL, "")]]

    total_pages = len(pages)

    for page_num, page_lines in enumerate(pages, start=1):
        parts: list[bytes] = [b"BT"]
        # Header
        parts.append(f"/F2 8 Tf".encode())
        header_text = f"{case_ref}  |  {classification}"
        parts.append(f"{_ML} {_PH - _MT} Td ({_esc_pdf(header_text)}) Tj".encode())
        # Page number
        pn_text = f"Page {page_num} of {total_pages}"
        pn_x = _PW - _MR - len(pn_text) * 5
        parts.append(f"{pn_x} {_PH - _MT} Td ({_esc_pdf(pn_text)}) Tj".encode())
        # Footer
        footer_text = "D-CIP Investigation Platform — Confidential"
        parts.append(f"{_ML} {_MB} Td ({_esc_pdf(footer_text)}) Tj".encode())
        # Reset position
        parts.append(f"{_ML} {TOP_Y} Td".encode())
        # Content
        current_size = _FONT_NORMAL
        parts.append(f"/F1 {current_size} Tf {int(line_height(current_size))} TL".encode())
        first_line = True
        for size, text in page_lines:
            if size != current_size:
                current_size = size
                # Start new position with new size
                parts.append(f"/F1 {size} Tf {int(line_height(size))} TL".encode())
            if first_line:
                first_line = False
            else:
                parts.append(b"T*")
            parts.append(f"({_esc_pdf(text)}) Tj".encode())
        parts.append(b"ET")

        stream = b"\n".join(parts)
        content_body = (
            f"<< /Length {len(stream)} >>\nstream\n".encode()
            + stream
            + b"\nendstream"
        )
        cid = add_obj(content_body)
        page_content_ids.append(cid)

    # Pages object
    pages_obj_offset = len(objects) + total_pages + 1
    real_page_ids: list[int] = []
    for cid in page_content_ids:
        body = (
            f"<< /Type /Page /Parent {pages_obj_offset} 0 R "
            f"/MediaBox [0 0 {_PW} {_PH}] "
            f"/Resources << /Font << /F1 {font_id} 0 R /F2 {bold_id} 0 R >> >> "
            f"/Contents {cid} 0 R >>"
        ).encode()
        real_page_ids.append(add_obj(body))

    kids = " ".join(f"{pid} 0 R" for pid in real_page_ids)
    pages_body = f"<< /Type /Pages /Kids [{kids}] /Count {total_pages} >>".encode()
    actual_pages_id = add_obj(pages_body)
    catalog_id = add_obj(f"<< /Type /Catalog /Pages {actual_pages_id} 0 R >>".encode())

    out = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets: list[int] = []
    for i, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"

    xref_pos = len(out)
    n = len(objects) + 1
    out += f"xref\n0 {n}\n".encode()
    out += b"0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += f"trailer\n<< /Size {n} /Root {catalog_id} 0 R >>\nstartxref\n{xref_pos}\n%%EOF".encode()
    return bytes(out)


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_docx(report: "InvestigationReport") -> bytes:
    """Generate a Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is not installed (pip install python-docx).")

    doc = Document()
    content = report.sections_content
    cover = content.get("cover", {})
    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")

    # Document properties
    doc.core_properties.title = report.title
    doc.core_properties.author = cover.get("generated_by", "D-CIP")
    doc.core_properties.description = f"Case report generated by D-CIP Investigation Platform"

    # Cover page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(cover.get("case_title") or report.title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 26, 46)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        f"Reference: {cover.get('case_reference', '')}\n"
        f"Report Type: {report.report_type.replace('_', ' ').title()}\n"
        f"Version: {report.version}\n"
        f"Classification: {cover.get('classification', '')}\n"
        f"Generated: {now_str}"
    ).font.size = Pt(10)

    doc.add_page_break()

    for sec_key, sec in content.items():
        if not isinstance(sec, dict) or sec_key == "cover":
            continue
        sec_title = sec.get("title", sec_key)
        doc.add_heading(sec_title, level=1)

        if sec.get("is_ai_generated"):
            p = doc.add_paragraph()
            run = p.add_run("⚠ AI-Generated Content — must be independently verified by the lead investigator.")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(146, 64, 14)

        _docx_section(doc, sec)
        doc.add_page_break()

    # Footer note
    doc.add_heading("Document Information", level=2)
    doc.add_paragraph(f"Report SHA-256: {report.content_hash or 'N/A'}")
    doc.add_paragraph("This report was generated by the D-CIP Investigation Platform.")
    doc.add_paragraph("AI-generated content is clearly marked and must be independently verified.")

    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_section(doc, sec: dict) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn
    import io

    sec_type = sec.get("type", "")

    if sec_type == "executive_summary":
        if sec.get("summary_text"):
            doc.add_paragraph(sec["summary_text"])
        for heading, key in [("Key Findings", "key_findings"), ("Potential Leads", "potential_leads"), ("Open Questions", "open_questions")]:
            items = sec.get(key, [])
            if items:
                doc.add_heading(heading, level=2)
                for it in items:
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(str(it))

    elif sec_type == "case_overview":
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Field"
        hdr[1].text = "Value"
        fields = [
            ("Reference", sec.get("reference_number")),
            ("Status", sec.get("status")),
            ("Priority", sec.get("priority")),
            ("Category", sec.get("category")),
            ("Evidence Count", sec.get("evidence_count")),
            ("Task Count", sec.get("task_count")),
            ("Created", _fmt_ts(sec.get("created_at"))),
        ]
        for k, v in fields:
            if v is not None:
                row = table.add_row().cells
                row[0].text = k
                row[1].text = str(v)
        if sec.get("description"):
            doc.add_paragraph(sec["description"])

    elif sec_type == "evidence_inventory":
        items = sec.get("items", [])
        doc.add_paragraph(f"Total items: {len(items)}")
        if items:
            table = doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(["File", "Type", "Size", "Status", "Entities"]):
                hdr[i].text = h
            for item in items:
                row = table.add_row().cells
                row[0].text = item.get("filename", "")
                row[1].text = item.get("mime_type", "")
                row[2].text = str(item.get("file_size", ""))
                row[3].text = item.get("status", "")
                row[4].text = str(item.get("entity_count", 0))

    elif sec_type == "timeline":
        items = sec.get("items", [])
        doc.add_paragraph(f"Total events: {len(items)}")
        for item in items:
            ts = _fmt_ts(item.get("event_timestamp"))
            p = doc.add_paragraph()
            p.add_run(f"[{ts}] {item.get('title', '')}").bold = True
            doc.add_paragraph(
                f"Type: {item.get('event_type', '')} | Confidence: {int(item.get('confidence', 0) * 100)}%"
            ).paragraph_format.left_indent = Pt(20)
            if item.get("description"):
                doc.add_paragraph(item["description"]).paragraph_format.left_indent = Pt(20)

    elif sec_type == "entities":
        doc.add_paragraph(f"Total: {sec.get('total_count', 0)} entities")
        for etype, ents in sec.get("by_type", {}).items():
            if not ents:
                continue
            doc.add_heading(etype.replace("_", " ").title(), level=2)
            for ent in ents[:20]:
                doc.add_paragraph(
                    f"{ent['value']} ({int(ent.get('confidence', 1) * 100)}%)",
                    style="List Bullet",
                )

    elif sec_type == "ai_findings":
        disc = sec.get("disclaimer")
        if disc:
            doc.add_paragraph(disc)
        cs = sec.get("case_summary", {})
        if cs.get("summary_text"):
            doc.add_heading("Case Analysis", level=2)
            doc.add_paragraph(cs["summary_text"])
        for heading, key in [("Key Findings", "key_findings"), ("Potential Leads", "potential_leads")]:
            items = cs.get(key, [])
            if items:
                doc.add_heading(heading, level=3)
                for it in items:
                    doc.add_paragraph(str(it), style="List Bullet")
        for es in sec.get("evidence_summaries", []):
            doc.add_heading(f"Evidence: {es.get('filename', '')}", level=2)
            if es.get("summary_text"):
                doc.add_paragraph(es["summary_text"])

    elif sec_type == "notes_tasks":
        notes = sec.get("notes", [])
        tasks = sec.get("tasks", [])
        if notes:
            doc.add_heading(f"Notes ({len(notes)})", level=2)
            for n in notes:
                pin = " [PINNED]" if n.get("is_pinned") else ""
                p = doc.add_paragraph()
                p.add_run(f"{n['title']}{pin}").bold = True
                doc.add_paragraph(f"By: {n.get('created_by','')} on {_fmt_ts(n.get('created_at'))}")
                if n.get("content"):
                    doc.add_paragraph(n["content"])
        if tasks:
            doc.add_heading(f"Tasks ({len(tasks)})", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            for i, h in enumerate(["Title", "Status", "Priority", "Due"]):
                hdr[i].text = h
            for t in tasks:
                row = table.add_row().cells
                row[0].text = t.get("title", "")
                row[1].text = t.get("status", "")
                row[2].text = t.get("priority", "")
                row[3].text = _fmt_ts(t.get("due_date"))

    elif sec_type == "chain_of_custody":
        for item in sec.get("items", []):
            doc.add_heading(item.get("filename", ""), level=2)
            doc.add_paragraph(f"SHA-256: {item.get('sha256_hash', 'N/A')}")
            evts = item.get("events", [])
            if evts:
                table = doc.add_table(rows=1, cols=4)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                for i, h in enumerate(["Action", "Actor", "Description", "Time"]):
                    hdr[i].text = h
                for ev in evts:
                    row = table.add_row().cells
                    row[0].text = ev.get("action", "")
                    row[1].text = ev.get("actor", "")
                    row[2].text = ev.get("description", "")
                    row[3].text = _fmt_ts(ev.get("timestamp"))
