"""Text extraction service — pulls readable text from evidence files.

Each extractor handles a specific file category. Extraction never modifies the
original file — only reads are performed. OCR is a fallback for files with no
embedded text (handled separately in ocr.py).
"""

from __future__ import annotations

import email
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

_PDF_EXTS = {"pdf"}
_DOCX_EXTS = {"docx", "doc"}
_XLSX_EXTS = {"xlsx", "xls"}
_TEXT_EXTS = {"txt", "csv", "log", "md", "ini", "cfg", "yaml", "yml"}
_JSON_EXTS = {"json"}
_XML_EXTS = {"xml"}
_EML_EXTS = {"eml", "msg"}

_MAX_TEXT_BYTES = 200_000  # 200 KB text cap per evidence item


def extract_text(file_path: str | Path, mime_type: str, file_extension: str) -> str | None:
    """Route to the appropriate extractor. Returns raw text or None."""
    path = Path(file_path)
    ext = file_extension.lower().lstrip(".")

    if not path.exists():
        return None

    try:
        if ext in _PDF_EXTS or mime_type == "application/pdf":
            return _extract_pdf(path)
        if ext in _DOCX_EXTS or "wordprocessingml" in mime_type:
            return _extract_docx(path)
        if ext in _XLSX_EXTS or "spreadsheetml" in mime_type:
            return _extract_xlsx(path)
        if ext in _EML_EXTS or mime_type in ("message/rfc822",):
            return _extract_eml(path)
        if ext in _TEXT_EXTS or mime_type.startswith("text/"):
            return _extract_plaintext(path)
        if ext in _JSON_EXTS or mime_type == "application/json":
            return _extract_plaintext(path)
        if ext in _XML_EXTS or mime_type in ("application/xml", "text/xml"):
            return _extract_plaintext(path)
    except Exception as exc:
        logger.warning("Text extraction failed for %s: %s", file_path, exc)

    return None


def _extract_pdf(path: Path) -> str | None:
    import pypdf
    reader = pypdf.PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text.strip())
    return "\n\n".join(parts)[:_MAX_TEXT_BYTES] or None


def _extract_docx(path: Path) -> str | None:
    import docx
    doc = docx.Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)[:_MAX_TEXT_BYTES] or None


def _extract_xlsx(path: Path) -> str | None:
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts: list[str] = []
    for sheet in wb.worksheets:
        parts.append(f"=== Sheet: {sheet.title} ===")
        for row in sheet.iter_rows(values_only=True):
            row_text = "\t".join(str(cell) for cell in row if cell is not None)
            if row_text.strip():
                parts.append(row_text)
    return "\n".join(parts)[:_MAX_TEXT_BYTES] or None


def _extract_eml(path: Path) -> str | None:
    with open(path, "rb") as f:
        msg = email.message_from_bytes(f.read())
    parts: list[str] = []
    # Headers
    for header in ("From", "To", "Cc", "Subject", "Date"):
        val = msg.get(header)
        if val:
            parts.append(f"{header}: {val}")
    parts.append("")
    # Body
    if msg.is_multipart():
        for part in msg.walk():
            ct = part.get_content_type()
            if ct == "text/plain":
                payload = part.get_payload(decode=True)
                if payload:
                    _decode_payload(payload, parts)
    else:
        payload = msg.get_payload(decode=True)
        if payload:
            _decode_payload(payload, parts)
    return "\n".join(parts)[:_MAX_TEXT_BYTES] or None


def _decode_payload(payload: bytes, parts: list[str]) -> None:
    try:
        import chardet
        detected = chardet.detect(payload)
        enc = detected.get("encoding") or "utf-8"
    except ImportError:
        enc = "utf-8"
    try:
        parts.append(payload.decode(enc, errors="replace"))
    except Exception:
        parts.append(payload.decode("utf-8", errors="replace"))


def _extract_plaintext(path: Path) -> str | None:
    raw = path.read_bytes()
    try:
        import chardet
        detected = chardet.detect(raw[:4096])
        enc = detected.get("encoding") or "utf-8"
    except ImportError:
        enc = "utf-8"
    try:
        return raw[:_MAX_TEXT_BYTES].decode(enc, errors="replace") or None
    except Exception:
        return raw[:_MAX_TEXT_BYTES].decode("utf-8", errors="replace") or None


def extract_image_metadata(file_path: str | Path) -> dict:
    """Extract EXIF/GPS from images using Pillow. Never raises."""
    result: dict = {}
    try:
        from PIL import Image
        from PIL.ExifTags import TAGS, GPSTAGS
        img = Image.open(str(file_path))
        result["dimensions"] = list(img.size)
        result["mode"] = img.mode
        exif_data = img._getexif()  # type: ignore[attr-defined]
        if exif_data:
            exif: dict = {}
            for tag_id, value in exif_data.items():
                tag = TAGS.get(tag_id, str(tag_id))
                if isinstance(value, bytes):
                    try:
                        value = value.decode("utf-8", errors="replace")
                    except Exception:
                        value = value.hex()
                exif[tag] = value
            result["exif"] = {k: str(v)[:200] for k, v in exif.items()}
            # GPS
            if "GPSInfo" in exif:
                gps_raw = exif_data.get(next(
                    k for k, v in exif_data.items() if TAGS.get(k) == "GPSInfo"
                ), {})
                gps: dict = {}
                for key, val in gps_raw.items():
                    gps[GPSTAGS.get(key, key)] = val
                lat = _parse_gps_coord(gps.get("GPSLatitude"), gps.get("GPSLatitudeRef"))
                lon = _parse_gps_coord(gps.get("GPSLongitude"), gps.get("GPSLongitudeRef"))
                if lat is not None and lon is not None:
                    result["gps"] = {"lat": lat, "lon": lon}
    except Exception:
        pass
    return result


def _parse_gps_coord(coord: object, ref: object) -> float | None:
    try:
        d, m, s = coord  # type: ignore[misc]
        val = float(d) + float(m) / 60 + float(s) / 3600
        if ref in ("S", "W"):
            val = -val
        return round(val, 6)
    except Exception:
        return None
